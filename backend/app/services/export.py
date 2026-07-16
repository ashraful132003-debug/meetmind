"""Export a meeting as PDF or Word.

Both formats are generated server-side from the same decrypted data, so what you
download matches exactly what the app shows — no "print the page" screenshot that
loses the action items or renders the sidebar into the document.

Library choices, and why:

* **reportlab** for PDF — pure Python, no system dependencies. WeasyPrint or
  wkhtmltopdf would give nicer HTML-to-PDF fidelity but both need GTK/Qt native
  libraries, which cannot be installed on a 512MB free host without admin.
* **python-docx** for Word — ~250KB, pure Python, writes real .docx that Word and
  Google Docs open without complaint.

Neither needs a network call, so export works offline exactly like the rest of
the local-first path.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from datetime import datetime

log = logging.getLogger(__name__)

# Matches the app's accent so an exported document looks like it came from the
# same product, not a different one.
ACCENT = (0.39, 0.40, 0.95)      # #6366f1
INK = (0.06, 0.09, 0.16)         # #0f172a
MUTED = (0.42, 0.45, 0.50)       # #6b7280


@dataclass
class ExportData:
    """Everything an export needs, already decrypted by the caller.

    A plain dataclass rather than the ORM model on purpose: this module must not
    be able to touch the database or trigger a lazy load, and it must never see
    an encrypted column.
    """

    title: str
    created_at: datetime
    duration_seconds: float
    language: str | None
    owner_name: str
    summary: str | None
    topics: list[str]
    sentiment: str | None
    speakers: list[dict]          # display_name, talk_seconds, word_count, color
    action_items: list[dict]      # task, owner_label, due_text, priority, done
    transcript: list[dict] | None  # speaker_name, start_time, text


def _duration_label(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h}h {m}m"
    return f"{m}m {s}s" if m else f"{s}s"


def _ts(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _strip_markdown(text: str) -> list[tuple[str, str]]:
    """Turn the summary's Markdown subset into (kind, text) pairs.

    Kinds: 'h2', 'bullet', 'para'. Both exporters consume this, so the PDF and
    the Word file can never drift apart in how they interpret a summary.
    """
    out: list[tuple[str, str]] = []
    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            out.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            out.append(("h2", line[2:].strip()))
        elif re.match(r"^[-*]\s+", line):
            out.append(("bullet", re.sub(r"^[-*]\s+", "", line).strip()))
        else:
            out.append(("para", line))
    return out


def _clean_inline(text: str) -> str:
    """Drop inline Markdown markers. The exporters apply their own styling."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


# --- PDF ---------------------------------------------------------------------


def build_pdf(data: ExportData) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from xml.sax.saxutils import escape

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=data.title,
        author="MeetMind",
        subject="Meeting summary",
    )

    base = getSampleStyleSheet()
    accent = colors.Color(*ACCENT)
    ink = colors.Color(*INK)
    muted = colors.Color(*MUTED)

    s_title = ParagraphStyle(
        "t", parent=base["Title"], fontSize=20, leading=25, textColor=ink,
        alignment=TA_LEFT, spaceAfter=2,
    )
    s_meta = ParagraphStyle("m", parent=base["Normal"], fontSize=8.5, textColor=muted, leading=12)
    s_h2 = ParagraphStyle(
        "h", parent=base["Heading2"], fontSize=9, leading=12, textColor=accent,
        spaceBefore=13, spaceAfter=5, fontName="Helvetica-Bold",
    )
    s_body = ParagraphStyle("b", parent=base["Normal"], fontSize=9.5, leading=14, textColor=ink)
    s_bullet = ParagraphStyle("bu", parent=s_body, leftIndent=9, bulletIndent=2, spaceAfter=3)
    s_cell = ParagraphStyle("c", parent=base["Normal"], fontSize=8.5, leading=11.5, textColor=ink)
    s_small = ParagraphStyle("s", parent=base["Normal"], fontSize=7.5, textColor=muted, leading=10)

    story: list = []

    story.append(Paragraph(escape(data.title), s_title))
    meta = " &nbsp;·&nbsp; ".join(
        filter(
            None,
            [
                data.created_at.strftime("%d %b %Y, %H:%M"),
                _duration_label(data.duration_seconds),
                f"{len(data.speakers)} speaker{'s' if len(data.speakers) != 1 else ''}",
                (data.language or "").upper() or None,
                f"shared by {escape(data.owner_name)}",
            ],
        )
    )
    story.append(Paragraph(meta, s_meta))
    story.append(Spacer(1, 5))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.Color(0.89, 0.91, 0.94)))
    story.append(Spacer(1, 4))

    if data.topics:
        story.append(Paragraph("TOPICS", s_h2))
        story.append(Paragraph(escape(" · ".join(data.topics)), s_body))

    if data.summary:
        for kind, text in _strip_markdown(data.summary):
            text = escape(_clean_inline(text))
            if kind == "h2":
                story.append(Paragraph(text.upper(), s_h2))
            elif kind == "bullet":
                story.append(Paragraph(text, s_bullet, bulletText="•"))
            else:
                story.append(Paragraph(text, s_body))
                story.append(Spacer(1, 3))

    # --- Action items --------------------------------------------------------
    story.append(Paragraph("ACTION ITEMS", s_h2))
    if data.action_items:
        rows = [
            [
                Paragraph("<b>Task</b>", s_small),
                Paragraph("<b>Owner</b>", s_small),
                Paragraph("<b>Due</b>", s_small),
                Paragraph("<b>Priority</b>", s_small),
            ]
        ]
        for a in data.action_items:
            task = escape(a["task"])
            if a.get("done"):
                task = f"<strike>{task}</strike>"
            rows.append(
                [
                    Paragraph(task, s_cell),
                    Paragraph(escape(a["owner_label"]), s_cell),
                    Paragraph(escape(a.get("due_text") or "—"), s_cell),
                    Paragraph(escape((a.get("priority") or "medium").title()), s_cell),
                ]
            )
        table = Table(rows, colWidths=[86 * mm, 32 * mm, 26 * mm, 26 * mm], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.97, 0.98, 0.99)),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.Color(0.85, 0.87, 0.91)),
                    ("GRID", (0, 1), (-1, -1), 0.3, colors.Color(0.90, 0.92, 0.95)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
    else:
        story.append(Paragraph("No action items were identified in this meeting.", s_body))

    # --- Speakers ------------------------------------------------------------
    if data.speakers:
        story.append(Paragraph("SPEAKING TIME", s_h2))
        total = sum(s["talk_seconds"] for s in data.speakers) or 1
        rows = [[Paragraph("<b>Speaker</b>", s_small), Paragraph("<b>Time</b>", s_small),
                 Paragraph("<b>Share</b>", s_small), Paragraph("<b>Words</b>", s_small)]]
        for sp in data.speakers:
            rows.append(
                [
                    Paragraph(escape(sp["display_name"]), s_cell),
                    Paragraph(_ts(sp["talk_seconds"]), s_cell),
                    Paragraph(f"{sp['talk_seconds'] / total * 100:.0f}%", s_cell),
                    Paragraph(str(sp.get("word_count", 0)), s_cell),
                ]
            )
        table = Table(rows, colWidths=[86 * mm, 28 * mm, 28 * mm, 28 * mm], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.97, 0.98, 0.99)),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.Color(0.85, 0.87, 0.91)),
                    ("GRID", (0, 1), (-1, -1), 0.3, colors.Color(0.90, 0.92, 0.95)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)

    # --- Transcript ----------------------------------------------------------
    if data.transcript:
        story.append(PageBreak())
        story.append(Paragraph("FULL TRANSCRIPT", s_h2))
        s_line = ParagraphStyle("l", parent=s_body, fontSize=8.5, leading=12.5, spaceAfter=4)
        for seg in data.transcript:
            story.append(
                Paragraph(
                    f'<font color="#94a3b8">[{_ts(seg["start_time"])}]</font> '
                    f'<b>{escape(seg["speaker_name"])}:</b> {escape(seg["text"])}',
                    s_line,
                )
            )

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.Color(0.6, 0.63, 0.69))
        canvas.drawString(
            20 * mm, 10 * mm,
            "Generated by MeetMind — transcription and analysis run on the owner's own machine.",
        )
        canvas.drawRightString(A4[0] - 20 * mm, 10 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


# --- Word --------------------------------------------------------------------


def build_docx(data: ExportData) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.core_properties.title = data.title
    doc.core_properties.author = "MeetMind"
    doc.core_properties.comments = "Generated by MeetMind"

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    accent_rgb = RGBColor(0x63, 0x66, 0xF1)
    muted_rgb = RGBColor(0x6B, 0x72, 0x80)

    h = doc.add_paragraph()
    run = h.add_run(data.title)
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        " · ".join(
            filter(
                None,
                [
                    data.created_at.strftime("%d %b %Y, %H:%M"),
                    _duration_label(data.duration_seconds),
                    f"{len(data.speakers)} speaker{'s' if len(data.speakers) != 1 else ''}",
                    (data.language or "").upper() or None,
                    f"shared by {data.owner_name}",
                ],
            )
        )
    )
    meta_run.font.size = Pt(8.5)
    meta_run.font.color.rgb = muted_rgb

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = accent_rgb

    if data.topics:
        heading("Topics")
        doc.add_paragraph(" · ".join(data.topics))

    if data.summary:
        for kind, text in _strip_markdown(data.summary):
            text = _clean_inline(text)
            if kind == "h2":
                heading(text)
            elif kind == "bullet":
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text)

    heading("Action items")
    if data.action_items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, label in enumerate(("Task", "Owner", "Due", "Priority")):
            cell = table.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(8.5)
        for a in data.action_items:
            cells = table.add_row().cells
            task_run = cells[0].paragraphs[0].add_run(a["task"])
            task_run.font.size = Pt(9)
            if a.get("done"):
                # A ticked-off item stays in the document — the record is what was
                # committed to, not just what is outstanding.
                task_run.font.strike = True
            for idx, val in (
                (1, a["owner_label"]),
                (2, a.get("due_text") or "—"),
                (3, (a.get("priority") or "medium").title()),
            ):
                r = cells[idx].paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)
    else:
        doc.add_paragraph("No action items were identified in this meeting.")

    if data.speakers:
        heading("Speaking time")
        total = sum(s["talk_seconds"] for s in data.speakers) or 1
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, label in enumerate(("Speaker", "Time", "Share", "Words")):
            cell = table.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(8.5)
        for sp in data.speakers:
            cells = table.add_row().cells
            for idx, val in (
                (0, sp["display_name"]),
                (1, _ts(sp["talk_seconds"])),
                (2, f"{sp['talk_seconds'] / total * 100:.0f}%"),
                (3, str(sp.get("word_count", 0))),
            ):
                r = cells[idx].paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)

    if data.transcript:
        doc.add_page_break()
        heading("Full transcript")
        for seg in data.transcript:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            t = p.add_run(f"[{_ts(seg['start_time'])}] ")
            t.font.size = Pt(8)
            t.font.color.rgb = muted_rgb
            sp = p.add_run(f"{seg['speaker_name']}: ")
            sp.bold = True
            sp.font.size = Pt(9)
            body = p.add_run(seg["text"])
            body.font.size = Pt(9)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "Generated by MeetMind — transcription and analysis run on the owner's own machine."
    )
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = muted_rgb

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(title: str, extension: str) -> str:
    """A filename that survives every OS.

    The title comes from the user (or the LLM), so it can contain slashes, colons
    and quotes — all of which break either the filesystem or the
    Content-Disposition header.
    """
    cleaned = re.sub(r"[^\w\s-]", "", title).strip()
    cleaned = re.sub(r"[\s_]+", "-", cleaned).lower()[:60].strip("-")
    return f"{cleaned or 'meeting'}.{extension}"
